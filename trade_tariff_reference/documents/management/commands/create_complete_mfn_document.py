import tempfile

from django.core.management.base import BaseCommand
from django.db import transaction

from docx import Document

from docxcompose.composer import Composer

from trade_tariff_reference.documents.mfn.constants import CLASSIFICATION, SCHEDULE
from trade_tariff_reference.schedule.models import Chapter, MFNDocument
from trade_tariff_reference.documents.utils import upload_generic_document_to_s3


class Command(BaseCommand):

    help = 'Command to combine all schedule/classification chapters in to one document'

    def add_arguments(self, parser):
        parser.add_argument('document_type', type=str, choices=[SCHEDULE, CLASSIFICATION])

    def handle(self, *args, **options):
        self.create_document(options['document_type'])

    def create_document(self, document_type):
        document = MFNDocument.objects.filter(document_type=document_type).first()
        document_field = f"{document_type}_document"

        if self.should_update_document(document, document_field):
            with tempfile.TemporaryDirectory(prefix='mfn_complete') as tmp_model_dir:
                document_name = self.generate_document(document_type, document_field, tmp_model_dir)
                self.update_database(document, document_type, document_name)

        else:
            self.stdout.write('No update found')

    def should_update_document(self, document, document_field):
        hashes = [getattr(chapter, f'{document_field}_check_sum') for chapter in Chapter.objects.all()]
        if document and set(hashes) == set(document.document_check_sum_list):
            return False
        return True

    @transaction.atomic
    def update_database(self, document, document_type, document_name):
        if not document:
            document = MFNDocument(document_type=document_type)

        hashes = [getattr(chapter, f'schedule_document_check_sum') for chapter in Chapter.objects.all()]
        document.document_check_sum_list = hashes
        upload_generic_document_to_s3(document, 'document', document_name, 'schedule.docx')

    def generate_document(self, document_type, document_field, tmp_dir):
        document_name = f'{tmp_dir}/{document_type}.docx'
        chapters = Chapter.objects.filter(id__in=[1, 2])
        first_chapter = chapters.first()
        document = getattr(first_chapter, document_field)
        master_document = Document(document)
        complete_document = Composer(master_document)
        for chapter in chapters.exclude(id=first_chapter.id):
            document = getattr(chapter, document_field)
            docx = Document(document)
            complete_document.append(docx)
            complete_document.save(document_name)
        return document_name
