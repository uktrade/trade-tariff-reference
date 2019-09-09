import logging
from html.parser import HTMLParser

from docx import Document

logger = logging.getLogger(__name__)


HEADER = (
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
    'xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main" '
    'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
    'xmlns:mv="urn:schemas-microsoft-com:mac:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
    'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
    'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
    'xmlns:w10="urn:schemas-microsoft-com:office:word" '
    'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
    'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
    'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
    'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
)


def update_description(html_code):
    document = Document()
    document.settings.element.remove_all()
    document_html_parser = DocumentHTMLParser(document)
    document_html_parser.add_paragraph_and_feed(html_code)
    xml = document_html_parser.paragraph._element.xml
    xml = xml.replace(HEADER, '<w:p>')
    return xml


class DocumentHTMLParser(HTMLParser):

    def __init__(self, document):
        HTMLParser.__init__(self)
        self.document = document

    def add_paragraph_and_feed(self, html):
        self.paragraph = self.document.add_paragraph()
        self.run = self.paragraph.add_run()
        self.feed(html)

    def handle_starttag(self, tag, attrs):
        self.run = self.paragraph.add_run()
        if tag == "i":
            self.run.italic = True
        if tag == "b":
            self.run.bold = True
        if tag == "u":
            self.run.underline = True
        if tag == "br":
            self.run.add_break()
        if tag in "li":
            self.run.add_text(u'- ')
        if tag == 'sub':
            self.run.font.subscript = True
        if tag == 'sup':
            self.run.font.superscript = True

    def handle_endtag(self, tag):
        if tag in ["br", "li", "ul", "ol"]:
            self.run.add_break()
        self.run = self.paragraph.add_run()

    def handle_data(self, data):
        self.run.add_text(data)
